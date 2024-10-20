from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.paginator import Paginator
from django.views import View
from django.utils.timezone import now
from django.http import StreamingHttpResponse
from docx import Document
import io

from .models import LogBook

# Create your views here.


@login_required(login_url="/auth/login")
def home(request):
    logs = LogBook.objects.filter(student=request.user)

    paginator = Paginator(logs, 5)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)

    context = {
        "logs": logs,
        "page_obj": page_obj,
    }
    return render(request, "student/home.html", context)


@login_required(login_url="/auth/login")
def add_log(request):
    try:
        print(now().date(), "current_Date")
        current_log = LogBook.objects.filter(date=now().date())
        if current_log:
            print(current_log, "current")
            messages.error(request, "Sorry A Log For Today Already Exists!")
            return redirect("student_home")
        if request.method == "POST":
            current_log = LogBook.objects.filter(date=now)
            print(current_log, "current")
            data = request.POST
            context = {"fieldValues": data}
            title = data.get("log_title")
            desc = data.get("log_desc")
            # date = data["log_date"]
            image = data.get("log_img")

            if not desc:
                messages.error(request, "Description Is Required")
                return render(request, "student/add_log.html", context)

            LogBook.objects.create(
                title=title, desc=desc, image=image, student=request.user
            )
            messages.success(request, "Log saved successfully")
            return redirect("student_home")
        return render(request, "student/add_log.html")

    except Exception as e:
        # log = LogBook.objects.get(pk=id)
        # context = {
        #     "fieldValues": log,
        # }
        print(repr(e), "error occured")

        if str(e).startswith("value too long for type character varying(20)"):
            messages.error(request, "Title field should be at most 20 characters")
            return render(
                request,
                "student/add_log.html",
            )

        messages.error(request, "Sorry An Error Occured")
        return render(
            request,
            "student/add_log.html",
        )

    # student_logs = LogBook.objects.filter(owner=request.user)


@login_required(login_url="/auth/login")
def edit_log(request, id):
    try:
        log = LogBook.objects.get(pk=id)
        print(log.image, "logdetails")

        context = {
            "fieldValues": log,
            "log": log,
        }

        if request.method == "POST":
            data = request.POST
            title = data.get("log_title")
            desc = data.get("log_desc")
            img = data.get("log_img")
            if not desc:
                messages.error(request, "Description Is Required")
                return render(request, "student/edit_log.html", context)
            log.title = title
            log.desc = desc
            log.student = request.user
            log.image = img

            log.save()
            messages.success(request, "Log edited successfully")
            return redirect("student_home")

        return render(request, "student/edit_log.html", context)
    except Exception as e:
        log = LogBook.objects.get(pk=id)
        context = {
            "fieldValues": log,
            "log": log,
        }
        print(repr(e), "error occured")

        if str(e).startswith("value too long for type character varying(20)"):
            messages.error(request, "Title field should be at most 20 characters")
            return render(request, "student/edit_log.html", context)

        messages.error(request, "Sorry An Error Occured")
        return render(request, "student/edit_log.html", context)


def show_log(request, id):
    print(id)
    log = LogBook.objects.get(pk=id)
    context = {
        "log": log,
    }

    return render(request, "student/show_log.html", context)


class Export_Word_View(View):
    def get(self, request, id):
        logs = LogBook.objects.filter(student=request.user)
        current_log = logs.get(pk=id)
        print(current_log, "logs1")

        # create an empty document object
        document = self.write_docx(current_log)

        # save document info
        buffer = io.BytesIO()
        document.save(buffer)  # save your memory stream
        buffer.seek(0)  # rewind the stream

        # put them to streaming content response
        # within docx content_type
        response = StreamingHttpResponse(
            streaming_content=buffer,  # use the stream's content
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        response["Content-Disposition"] = (
            "attachment;filename=logbook_day_" + str(current_log.date) + ".docx"
        )
        response["Content-Encoding"] = "UTF-8"

        return response

    def write_docx(self, current_log: LogBook):
        document = Document()

        document.add_heading(current_log.title, 0)

        document.add_paragraph(current_log.desc)

        return document


export_word_view = Export_Word_View.as_view()
